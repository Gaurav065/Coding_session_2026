import docx
import re
import sys

def clean_docx(input_path, output_path):
    print(f"Loading {input_path}...")
    try:
        doc = docx.Document(input_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    # Clean paragraphs
    for p in doc.paragraphs:
        # Check if paragraph only contains whitespace
        text = p.text
        if not text.strip():
            # Check if it has runs with images or page breaks before deleting
            has_special = False
            for run in p.runs:
                # Check for drawing/images
                if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                    has_special = True
                    break
                # Check for page breaks (w:br with type="page")
                if run._element.xpath('.//w:br[@w:type="page"]'):
                    # It's a manual page break. If the user wants to skip empty pages,
                    # we might want to remove this if it results in an empty page, 
                    # but it's tricky. Let's keep it to be safe, or actually remove 
                    # multiple consecutive page breaks.
                    pass
            
            if not has_special:
                # Remove the empty paragraph
                p._element.getparent().remove(p._element)
                p._p = p._element = None
                continue

        # If it has text, clean extra spaces
        if text.strip():
            # Clean text in each run to preserve formatting
            # This is tricky because replacing across runs can break things.
            # A simpler way is to replace extra spaces within each run.
            for run in p.runs:
                if run.text:
                    # Replace multiple spaces with a single space
                    run.text = re.sub(r' {2,}', ' ', run.text)

    # Clean tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if text.strip():
                        for run in p.runs:
                            if run.text:
                                run.text = re.sub(r' {2,}', ' ', run.text)

    print(f"Saving to {output_path}...")
    doc.save(output_path)
    print("Done!")

if __name__ == "__main__":
    input_file = r"C:\Users\GauravPatel\Downloads\Metadata_Control_Application_Solution_Documentation (1).docx"
    output_file = r"C:\Users\GauravPatel\Downloads\Metadata_Control_Application_Solution_Documentation_Cleaned.docx"
    clean_docx(input_file, output_file)
